"""
PDF and DOCX report generation for artwork analysis.
"""
from datetime import datetime
from pathlib import Path
from typing import Literal, Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from backend.core.config import settings
from backend.core.logging import log
from backend.core.models import AnalysisResult


class ReportBuilder:
    """
    Generates professional PDF and DOCX reports for artwork analysis.
    """

    def __init__(self, output_dir: Optional[Path] = None):
        """
        Initialize report builder.

        Args:
            output_dir: Directory for saving reports
        """
        self.output_dir = output_dir or Path(settings.REPORT_OUTPUT_DIR)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_report(
        self,
        analysis_result: AnalysisResult,
        output_path: Optional[str] = None,
        format: Literal["pdf", "docx"] = "pdf",
        include_image: bool = True
    ) -> str:
        """
        Generate a complete artwork analysis report.

        Args:
            analysis_result: Complete analysis result
            output_path: Custom output path (optional)
            format: Report format ('pdf' or 'docx')
            include_image: Whether to include the artwork image

        Returns:
            Path to generated report file
        """
        log.info(f"Generating {format.upper()} report...")

        if format == "pdf":
            return self._generate_pdf(analysis_result, output_path, include_image)
        elif format == "docx":
            return self._generate_docx(analysis_result, output_path, include_image)
        else:
            raise ValueError(f"Unsupported format: {format}")

    def _generate_pdf(
        self,
        analysis_result: AnalysisResult,
        output_path: Optional[str],
        include_image: bool
    ) -> str:
        """Generate PDF report using ReportLab."""
        # Determine output path
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"artwork_analysis_{timestamp}.pdf"
        else:
            output_path = Path(output_path)

        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18,
        )

        # Container for PDF elements
        elements = []

        # Styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#2C3E50'),
            spaceAfter=30,
            alignment=1  # Center
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#34495E'),
            spaceAfter=12,
            spaceBefore=12
        )

        # Title
        elements.append(Paragraph("Artwork Analysis Report", title_style))
        elements.append(Spacer(1, 0.2 * inch))

        # Metadata
        metadata_text = f"<b>Analysis Date:</b> {analysis_result.metadata.created_at.strftime('%Y-%m-%d %H:%M:%S')}<br/>"
        metadata_text += f"<b>Source:</b> {analysis_result.metadata.source}<br/>"
        if analysis_result.metadata.notes:
            metadata_text += f"<b>Notes:</b> {analysis_result.metadata.notes}<br/>"

        elements.append(Paragraph(metadata_text, styles['Normal']))
        elements.append(Spacer(1, 0.3 * inch))

        # Image (if available and requested)
        if include_image and Path(analysis_result.image_path).exists():
            try:
                img = Image(analysis_result.image_path, width=4*inch, height=3*inch)
                elements.append(img)
                elements.append(Spacer(1, 0.3 * inch))
            except Exception as e:
                log.warning(f"Could not include image in report: {e}")

        # Style & Epoch
        elements.append(Paragraph("Style & Epoch Analysis", heading_style))
        style_data = [
            ["Epoch", analysis_result.style_estimation.epoch],
            ["Style", analysis_result.style_estimation.style],
            ["Confidence", f"{analysis_result.style_estimation.confidence:.1%}"]
        ]
        style_table = Table(style_data, colWidths=[2*inch, 3.5*inch])
        style_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ECF0F1')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey)
        ]))
        elements.append(style_table)
        elements.append(Spacer(1, 0.2 * inch))

        # Artist Candidates
        elements.append(Paragraph("Artist Matching", heading_style))
        if analysis_result.artist_candidates:
            artist_data = [["Artist", "Similarity", "Epoch/Style"]]
            for candidate in analysis_result.artist_candidates[:5]:
                artist_data.append([
                    candidate.name,
                    f"{candidate.similarity:.1%}",
                    f"{candidate.epoch or ''} / {candidate.style or ''}"
                ])

            artist_table = Table(artist_data, colWidths=[2.5*inch, 1.5*inch, 1.5*inch])
            artist_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey)
            ]))
            elements.append(artist_table)
        else:
            elements.append(Paragraph("No artist matches found.", styles['Normal']))

        elements.append(Spacer(1, 0.2 * inch))

        # Authenticity
        elements.append(Paragraph("Authenticity Assessment", heading_style))
        auth_score_color = self._get_score_color(analysis_result.authenticity_score)
        auth_text = f"<b>Overall Score:</b> <font color='{auth_score_color}'>{analysis_result.authenticity_score}/100</font>"
        elements.append(Paragraph(auth_text, styles['Normal']))
        elements.append(Spacer(1, 0.1 * inch))

        # ... (continue with condition, valuation, and provenance sections)
        # Condition
        elements.append(Paragraph("Condition Analysis", heading_style))
        condition = analysis_result.condition
        condition_data = [
            ["Craquelé", "Yes" if condition.craquele else "No"],
            ["Yellowing", "Yes" if condition.yellowing else "No"],
            ["Stains", "Yes" if condition.stains else "No"],
            ["Damage Score", f"{condition.damage_score:.1%}"],
            ["Notes", condition.notes]
        ]
        condition_table = Table(condition_data, colWidths=[2*inch, 3.5*inch])
        condition_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ECF0F1')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))
        elements.append(condition_table)
        elements.append(Spacer(1, 0.2 * inch))

        # Valuation
        elements.append(Paragraph("Market Valuation", heading_style))
        valuation = analysis_result.valuation
        val_text = f"<b>Estimated Value:</b> €{valuation.estimated_value:,.0f}<br/>"
        val_text += f"<b>Range:</b> €{valuation.min:,.0f} - €{valuation.max:,.0f}<br/>"
        val_text += f"<b>Confidence:</b> {valuation.confidence.upper()}<br/>"
        val_text += f"<b>Rationale:</b> {valuation.rationale}"
        elements.append(Paragraph(val_text, styles['Normal']))
        elements.append(Spacer(1, 0.2 * inch))

        # Provenance
        elements.append(Paragraph("Provenance Research", heading_style))
        prov = analysis_result.provenance
        prov_text = f"<b>Reverse Image Hits:</b> {prov.reverse_image_hits}<br/>"
        prov_text += f"<b>Auction Records:</b> {len(prov.auction_history)}<br/>"
        prov_text += f"<b>Notes:</b> {prov.notes}"
        elements.append(Paragraph(prov_text, styles['Normal']))

        # Footer
        elements.append(Spacer(1, 0.5 * inch))
        footer_text = "<i>This report was generated automatically by Gemäldeagent. "
        footer_text += "Professional expert appraisal is recommended for verification.</i>"
        elements.append(Paragraph(footer_text, styles['Italic']))

        # Build PDF
        doc.build(elements)

        log.info(f"PDF report generated: {output_path}")
        return str(output_path)

    def _generate_docx(
        self,
        analysis_result: AnalysisResult,
        output_path: Optional[str],
        include_image: bool
    ) -> str:
        """Generate DOCX report using python-docx."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Determine output path
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"artwork_analysis_{timestamp}.docx"
        else:
            output_path = Path(output_path)

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading('Artwork Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        doc.add_paragraph(
            f"Analysis Date: {analysis_result.metadata.created_at.strftime('%Y-%m-%d %H:%M:%S')}"
        )
        doc.add_paragraph(f"Source: {analysis_result.metadata.source}")
        if analysis_result.metadata.notes:
            doc.add_paragraph(f"Notes: {analysis_result.metadata.notes}")

        doc.add_paragraph()  # Spacer

        # Image
        if include_image and Path(analysis_result.image_path).exists():
            try:
                doc.add_picture(analysis_result.image_path, width=Inches(4))
                doc.add_paragraph()
            except Exception as e:
                log.warning(f"Could not include image in DOCX: {e}")

        # Style & Epoch
        doc.add_heading('Style & Epoch Analysis', 1)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.rows[0].cells[0].text = 'Epoch'
        table.rows[0].cells[1].text = analysis_result.style_estimation.epoch
        table.rows[1].cells[0].text = 'Style'
        table.rows[1].cells[1].text = analysis_result.style_estimation.style
        table.rows[2].cells[0].text = 'Confidence'
        table.rows[2].cells[1].text = f"{analysis_result.style_estimation.confidence:.1%}"

        # Artist Matching
        doc.add_heading('Artist Matching', 1)
        if analysis_result.artist_candidates:
            artist_table = doc.add_table(rows=len(analysis_result.artist_candidates[:5])+1, cols=3)
            artist_table.style = 'Light Grid Accent 1'
            artist_table.rows[0].cells[0].text = 'Artist'
            artist_table.rows[0].cells[1].text = 'Similarity'
            artist_table.rows[0].cells[2].text = 'Epoch/Style'

            for i, candidate in enumerate(analysis_result.artist_candidates[:5], 1):
                artist_table.rows[i].cells[0].text = candidate.name
                artist_table.rows[i].cells[1].text = f"{candidate.similarity:.1%}"
                artist_table.rows[i].cells[2].text = f"{candidate.epoch or ''} / {candidate.style or ''}"

        # Authenticity
        doc.add_heading('Authenticity Assessment', 1)
        doc.add_paragraph(f"Overall Score: {analysis_result.authenticity_score}/100")

        # Condition
        doc.add_heading('Condition Analysis', 1)
        condition = analysis_result.condition
        doc.add_paragraph(f"Craquelé: {'Yes' if condition.craquele else 'No'}")
        doc.add_paragraph(f"Yellowing: {'Yes' if condition.yellowing else 'No'}")
        doc.add_paragraph(f"Stains: {'Yes' if condition.stains else 'No'}")
        doc.add_paragraph(f"Damage Score: {condition.damage_score:.1%}")
        doc.add_paragraph(f"Notes: {condition.notes}")

        # Valuation
        doc.add_heading('Market Valuation', 1)
        valuation = analysis_result.valuation
        doc.add_paragraph(f"Estimated Value: €{valuation.estimated_value:,.0f}")
        doc.add_paragraph(f"Range: €{valuation.min:,.0f} - €{valuation.max:,.0f}")
        doc.add_paragraph(f"Confidence: {valuation.confidence.upper()}")
        doc.add_paragraph(f"Rationale: {valuation.rationale}")

        # Provenance
        doc.add_heading('Provenance Research', 1)
        prov = analysis_result.provenance
        doc.add_paragraph(f"Reverse Image Hits: {prov.reverse_image_hits}")
        doc.add_paragraph(f"Auction Records: {len(prov.auction_history)}")
        doc.add_paragraph(f"Notes: {prov.notes}")

        # Footer
        footer_para = doc.add_paragraph()
        footer_para.add_run(
            "This report was generated automatically by Gemäldeagent. "
            "Professional expert appraisal is recommended for verification."
        ).italic = True

        # Save
        doc.save(str(output_path))

        log.info(f"DOCX report generated: {output_path}")
        return str(output_path)

    def _get_score_color(self, score: int) -> str:
        """Get color code based on score."""
        if score >= 80:
            return '#27AE60'  # Green
        elif score >= 60:
            return '#F39C12'  # Orange
        else:
            return '#E74C3C'  # Red


# Global report builder instance
report_builder = ReportBuilder()
